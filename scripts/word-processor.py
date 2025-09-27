#!/usr/bin/env python3
"""
Word Document Processor for Cursor
==================================

A comprehensive tool for extracting, analyzing, and working with Word documents
directly in Cursor. This fills a critical gap in the Cursor ecosystem.

Features:
- Extract text from .docx files with formatting preservation
- Convert to markdown for better Cursor integration
- Batch process multiple Word documents
- Extract metadata (author, creation date, etc.)
- Generate summaries and analysis
- Cursor-friendly output formatting

Usage:
    python word-processor.py extract <file.docx> [--output <output.md>]
    python word-processor.py batch <directory> [--output <output_dir>]
    python word-processor.py analyze <file.docx> [--summary]
    python word-processor.py convert <file.docx> [--format markdown|text|json]
"""

import argparse
import os
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import re

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

class WordProcessor:
    """Main class for processing Word documents."""
    
    def __init__(self):
        self.supported_formats = ['.docx']
        self.output_formats = ['markdown', 'text', 'json']
    
    def extract_text(self, file_path: str, preserve_formatting: bool = True) -> Dict:
        """
        Extract text and metadata from a Word document.
        
        Args:
            file_path: Path to the Word document
            preserve_formatting: Whether to preserve basic formatting (bold, italic, etc.)
            
        Returns:
            Dictionary containing text, metadata, and formatting information
        """
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract text content
            if preserve_formatting:
                content = self._extract_formatted_text(doc)
            else:
                content = self._extract_plain_text(doc)
            
            # Extract structure (headings, lists, etc.)
            structure = self._extract_structure(doc)
            
            return {
                'file_path': file_path,
                'metadata': metadata,
                'content': content,
                'structure': structure,
                'extraction_date': datetime.now().isoformat(),
                'word_count': len(content.split()),
                'character_count': len(content)
            }
            
        except Exception as e:
            return {
                'error': f"Failed to process {file_path}: {str(e)}",
                'file_path': file_path
            }
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """Extract document metadata."""
        metadata = {}
        
        # Core properties
        core_props = doc.core_properties
        metadata.update({
            'title': core_props.title or 'Untitled',
            'author': core_props.author or 'Unknown',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'last_modified_by': core_props.last_modified_by or '',
            'revision': core_props.revision or 0,
            'version': core_props.version or ''
        })
        
        return metadata
    
    def _extract_plain_text(self, doc: Document) -> str:
        """Extract plain text without formatting."""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return '\n\n'.join(text_parts)
    
    def _extract_formatted_text(self, doc: Document) -> str:
        """Extract text with basic markdown formatting."""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                text_parts.append('')
                continue
                
            # Check if it's a heading
            if self._is_heading(paragraph):
                level = self._get_heading_level(paragraph)
                text_parts.append(f"{'#' * level} {paragraph.text.strip()}")
            else:
                # Process runs for formatting
                formatted_text = self._format_paragraph(paragraph)
                text_parts.append(formatted_text)
        
        return '\n\n'.join(text_parts)
    
    def _is_heading(self, paragraph) -> bool:
        """Check if paragraph is a heading."""
        return paragraph.style.name.startswith('Heading')
    
    def _get_heading_level(self, paragraph) -> int:
        """Get heading level (1-6)."""
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 1
        return 1
    
    def _format_paragraph(self, paragraph) -> str:
        """Format a paragraph with basic markdown."""
        formatted_runs = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
                
            # Apply formatting
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            
            formatted_runs.append(text)
        
        return ''.join(formatted_runs)
    
    def _extract_structure(self, doc: Document) -> Dict:
        """Extract document structure (headings, lists, etc.)."""
        structure = {
            'headings': [],
            'paragraphs': [],
            'tables': [],
            'lists': []
        }
        
        for i, paragraph in enumerate(doc.paragraphs):
            if self._is_heading(paragraph):
                level = self._get_heading_level(paragraph)
                structure['headings'].append({
                    'level': level,
                    'text': paragraph.text.strip(),
                    'position': i
                })
            else:
                structure['paragraphs'].append({
                    'text': paragraph.text.strip(),
                    'position': i,
                    'style': paragraph.style.name
                })
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            structure['tables'].append({
                'index': i,
                'data': table_data,
                'rows': len(table_data),
                'cols': len(table_data[0]) if table_data else 0
            })
        
        return structure
    
    def convert_to_markdown(self, file_path: str, output_path: Optional[str] = None) -> str:
        """Convert Word document to markdown."""
        result = self.extract_text(file_path, preserve_formatting=True)
        
        if 'error' in result:
            return f"Error: {result['error']}"
        
        # Create markdown content
        md_content = []
        
        # Add metadata header
        metadata = result['metadata']
        md_content.append(f"# {metadata['title']}")
        md_content.append("")
        md_content.append(f"**Author:** {metadata['author']}")
        md_content.append(f"**Created:** {metadata['created'] or 'Unknown'}")
        md_content.append(f"**Modified:** {metadata['modified'] or 'Unknown'}")
        md_content.append(f"**Word Count:** {result['word_count']}")
        md_content.append("")
        md_content.append("---")
        md_content.append("")
        
        # Add content
        md_content.append(result['content'])
        
        # Add structure summary
        structure = result['structure']
        if structure['headings']:
            md_content.append("")
            md_content.append("## Document Structure")
            md_content.append("")
            for heading in structure['headings']:
                indent = "  " * (heading['level'] - 1)
                md_content.append(f"{indent}- {heading['text']}")
        
        markdown_text = '\n'.join(md_content)
        
        # Save to file if output path specified
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_text)
            print(f"Markdown saved to: {output_path}")
        
        return markdown_text
    
    def batch_process(self, directory: str, output_dir: Optional[str] = None) -> List[Dict]:
        """Process all Word documents in a directory."""
        directory = Path(directory)
        if not directory.exists():
            return [{'error': f"Directory not found: {directory}"}]
        
        if output_dir:
            output_dir = Path(output_dir)
            output_dir.mkdir(exist_ok=True)
        
        results = []
        word_files = list(directory.glob('*.docx'))
        
        print(f"Found {len(word_files)} Word documents to process...")
        
        for file_path in word_files:
            print(f"Processing: {file_path.name}")
            
            result = self.extract_text(str(file_path))
            results.append(result)
            
            # Save markdown if output directory specified
            if output_dir and 'error' not in result:
                md_file = output_dir / f"{file_path.stem}.md"
                self.convert_to_markdown(str(file_path), str(md_file))
        
        return results
    
    def analyze_document(self, file_path: str, include_summary: bool = True) -> Dict:
        """Analyze document content and structure."""
        result = self.extract_text(file_path)
        
        if 'error' in result:
            return result
        
        analysis = {
            'file_info': {
                'name': Path(file_path).name,
                'size': os.path.getsize(file_path),
                'word_count': result['word_count'],
                'character_count': result['character_count']
            },
            'metadata': result['metadata'],
            'structure': result['structure']
        }
        
        if include_summary:
            analysis['summary'] = self._generate_summary(result)
        
        return analysis
    
    def _generate_summary(self, result: Dict) -> Dict:
        """Generate document summary."""
        content = result['content']
        structure = result['structure']
        
        # Extract key topics from headings
        topics = [h['text'] for h in structure['headings'] if h['level'] <= 2]
        
        # Basic readability metrics
        sentences = len(re.split(r'[.!?]+', content))
        avg_words_per_sentence = result['word_count'] / max(sentences, 1)
        
        return {
            'main_topics': topics[:5],  # Top 5 topics
            'total_headings': len(structure['headings']),
            'total_paragraphs': len(structure['paragraphs']),
            'total_tables': len(structure['tables']),
            'readability': {
                'sentences': sentences,
                'avg_words_per_sentence': round(avg_words_per_sentence, 1)
            }
        }

def main():
    """Main CLI interface."""
    parser = argparse.ArgumentParser(
        description="Word Document Processor for Cursor",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Extract text from a Word document
  python word-processor.py extract document.docx
  
  # Convert to markdown
  python word-processor.py convert document.docx --format markdown --output document.md
  
  # Batch process directory
  python word-processor.py batch ./documents --output ./markdown
  
  # Analyze document
  python word-processor.py analyze document.docx --summary
        """
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Available commands')
    
    # Extract command
    extract_parser = subparsers.add_parser('extract', help='Extract text from Word document')
    extract_parser.add_argument('file', help='Path to Word document')
    extract_parser.add_argument('--output', '-o', help='Output file path')
    extract_parser.add_argument('--format', choices=['text', 'markdown', 'json'], 
                               default='markdown', help='Output format')
    
    # Convert command
    convert_parser = subparsers.add_parser('convert', help='Convert Word document to other formats')
    convert_parser.add_argument('file', help='Path to Word document')
    convert_parser.add_argument('--output', '-o', help='Output file path')
    convert_parser.add_argument('--format', choices=['markdown', 'text', 'json'], 
                               default='markdown', help='Output format')
    
    # Batch command
    batch_parser = subparsers.add_parser('batch', help='Process multiple Word documents')
    batch_parser.add_argument('directory', help='Directory containing Word documents')
    batch_parser.add_argument('--output', '-o', help='Output directory for converted files')
    
    # Analyze command
    analyze_parser = subparsers.add_parser('analyze', help='Analyze document structure and content')
    analyze_parser.add_argument('file', help='Path to Word document')
    analyze_parser.add_argument('--summary', action='store_true', help='Include document summary')
    analyze_parser.add_argument('--output', '-o', help='Output file for analysis results')
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return
    
    processor = WordProcessor()
    
    try:
        if args.command == 'extract':
            result = processor.extract_text(args.file)
            
            if 'error' in result:
                print(f"Error: {result['error']}")
                return
            
            if args.format == 'json':
                output = json.dumps(result, indent=2, ensure_ascii=False)
            elif args.format == 'text':
                output = result['content']
            else:  # markdown
                output = processor.convert_to_markdown(args.file)
            
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(output)
                print(f"Output saved to: {args.output}")
            else:
                print(output)
        
        elif args.command == 'convert':
            if args.format == 'markdown':
                output = processor.convert_to_markdown(args.file, args.output)
                if not args.output:
                    print(output)
            else:
                result = processor.extract_text(args.file)
                if 'error' in result:
                    print(f"Error: {result['error']}")
                    return
                
                if args.format == 'json':
                    output = json.dumps(result, indent=2, ensure_ascii=False)
                else:  # text
                    output = result['content']
                
                if args.output:
                    with open(args.output, 'w', encoding='utf-8') as f:
                        f.write(output)
                    print(f"Output saved to: {args.output}")
                else:
                    print(output)
        
        elif args.command == 'batch':
            results = processor.batch_process(args.directory, args.output)
            
            # Print summary
            successful = len([r for r in results if 'error' not in r])
            failed = len(results) - successful
            
            print(f"\nBatch processing complete:")
            print(f"  Successfully processed: {successful}")
            print(f"  Failed: {failed}")
            
            if failed > 0:
                print("\nFailed files:")
                for result in results:
                    if 'error' in result:
                        print(f"  - {result['file_path']}: {result['error']}")
        
        elif args.command == 'analyze':
            result = processor.analyze_document(args.file, args.summary)
            
            if 'error' in result:
                print(f"Error: {result['error']}")
                return
            
            output = json.dumps(result, indent=2, ensure_ascii=False)
            
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    f.write(output)
                print(f"Analysis saved to: {args.output}")
            else:
                print(output)
    
    except KeyboardInterrupt:
        print("\nOperation cancelled by user.")
    except Exception as e:
        print(f"Unexpected error: {e}")

if __name__ == '__main__':
    main()
